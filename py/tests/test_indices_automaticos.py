"""Índices automáticos en el .docx: contenido, tablas e ilustraciones.

Requisito de tesis: los índices NO pueden escribirse a mano. Tienen que ser
campos de Word (TOC) que Word y OnlyOffice actualicen solos, y las leyendas
tienen que numerarse con campos SEQ para que el índice las recoja.
"""
from pathlib import Path

import pytest
from docx import Document

from dockit.generadores import docx as gd

GUION = {
    "tipo": "docx", "titulo": "El problema del oráculo", "autor": "Autor",
    "bloques": [
        {"clase": "indice"},
        {"clase": "indice_tablas"},
        {"clase": "indice_figuras"},
        {"clase": "titulo", "nivel": 1, "texto": "Introducción"},
        {"clase": "parrafo", "texto": "Texto de prueba."},
        {"clase": "tabla", "cabecera": ["A", "B"], "filas": [["1", "2"]],
         "leyenda": "Comparación de enfoques"},
        {"clase": "titulo", "nivel": 2, "texto": "Antecedentes"},
        {"clase": "figura", "ruta": "no-existe.png", "leyenda": "Esquema del flujo"},
    ],
}


def xml_de(destino) -> str:
    return Document(str(destino)).element.xml


@pytest.fixture
def doc(tmp_path):
    d = tmp_path / "t.docx"
    gd.generar(GUION, str(d), {}, {})
    return d


def test_indice_de_contenido_es_un_campo_toc(doc):
    x = xml_de(doc)
    assert "TOC" in x, "no hay ningún campo TOC en el documento"
    assert 'w:fldCharType="begin"' in x, "el campo no está bien formado"


def test_hay_tres_indices_distintos(doc):
    x = xml_de(doc)
    assert 'TOC \\o' in x or 'TOC \\\\o' in x, "falta el índice de contenido"
    assert 'Tabla' in x and 'Ilustración' in x, \
        "faltan los índices de tablas o de ilustraciones"


def test_las_leyendas_usan_campos_seq(doc):
    """Sin SEQ, la numeración es manual y el índice no las recoge."""
    x = xml_de(doc)
    assert "SEQ Tabla" in x, "la leyenda de tabla no numera con SEQ"
    assert "SEQ Ilustración" in x, "la leyenda de figura no numera con SEQ"


def test_las_leyendas_llevan_estilo_caption(doc):
    d = Document(str(doc))
    estilos = {p.style.name for p in d.paragraphs if p.text.strip()}
    assert "Caption" in estilos, \
        "sin el estilo Caption, el índice de tablas/figuras queda vacío"


def test_el_documento_pide_actualizar_campos_al_abrir(doc):
    """Si no, el usuario ve 'Actualice el índice' en vez del índice."""
    import zipfile
    with zipfile.ZipFile(doc) as z:
        assert "word/settings.xml" in z.namelist()
        settings = z.read("word/settings.xml").decode("utf-8")
    assert "updateFields" in settings, \
        "el documento no pide a Word que actualice los campos al abrirlo"


def test_los_titulos_usan_estilos_de_encabezado(doc):
    d = Document(str(doc))
    niveles = {p.style.name for p in d.paragraphs
               if p.style.name.startswith("Heading")}
    assert "Heading 1" in niveles and "Heading 2" in niveles, \
        "el TOC se construye desde los estilos Heading; sin ellos sale vacío"


def test_el_numero_de_pagina_es_un_campo_no_la_palabra_PAGE(doc):
    """Bug real: el pie mostraba «PAGE» literal. El campo estaba mal formado:
    se creaba un elemento <w:fldbegin> en vez de <w:fldChar w:fldCharType>."""
    import zipfile
    with zipfile.ZipFile(doc) as z:
        pies = [n for n in z.namelist() if n.startswith("word/footer")]
        assert pies, "no hay pie de página"
        xml = z.read(pies[0]).decode("utf-8")
    assert "fldbegin" not in xml, "elemento inventado: el campo no es válido"
    assert 'w:fldCharType="begin"' in xml, "el campo de página no está bien formado"
    assert "PAGE" in xml, "falta la instrucción PAGE"
