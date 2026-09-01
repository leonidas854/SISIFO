# -*- coding: utf-8 -*-
"""Documento académico sobre las seis guías del proceso forense digital."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x14, 0x2A, 0x45)
ACCENT = RGBColor(0x0E, 0x6C, 0x82)
GREY   = RGBColor(0x5A, 0x6A, 0x78)
BODY   = 'Calibri'

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = BODY
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), BODY)
pf = normal.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for name, size, color, before, after in [
    ('Heading 1', 16, NAVY, 18, 8),
    ('Heading 2', 13, ACCENT, 14, 6),
    ('Heading 3', 11.5, NAVY, 10, 4),
]:
    st = doc.styles[name]
    st.font.name = BODY; st.font.size = Pt(size)
    st.font.color.rgb = color; st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    st.paragraph_format.keep_with_next = True

# ── Utilidades ────────────────────────────────────────────────────────────
def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:color'), 'auto'); el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

def cell_text(cell, text, bold=False, size=9.5, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = BODY; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color

def fixed_layout(t, widths):
    t.autofit = False
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(Cm(w).twips)))

def repeat_header(t):
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)

def no_row_split(t):
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); trPr.insert(0, cs)

def make_table(headers, rows, widths, size=9.5, highlight=None, bold_first=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, '142A45')
        cell_text(c, h, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        hl = highlight is not None and ri == highlight
        for i, v in enumerate(row):
            if hl: shade(cells[i], 'DDEEF3')
            cell_text(cells[i], v, bold=(bold_first and i == 0) or (hl and i == 0), size=size)
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Cm(widths[i])
    fixed_layout(t, widths)
    repeat_header(t)
    no_row_split(t)
    return t

def para(text, italic=False, size=None, color=None, align=None, space_after=None, bold=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODY; r.italic = italic; r.bold = bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullets(items, style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.name = BODY
            r2 = p.add_run(' ' + it[1]); r2.font.name = BODY
        else:
            r = p.add_run(it); r.font.name = BODY

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = BODY; r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY

def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p.add_run('Guías del proceso forense digital  ·  Informática Forense  ·  Página ')
    r0.font.size = Pt(8); r0.font.name = BODY; r0.font.color.rgb = GREY
    r = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r._r.addnext(fld)
    rr = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '16'); rpr.append(sz); rr.append(rpr)
    t = OxmlElement('w:t'); t.text = '1'; rr.append(t); fld.append(rr)

# ══════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════
para('ESCUELA MILITAR DE INGENIERÍA', size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para('«Mcal. Antonio José de Sucre»', size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=4)
para('UNIDAD ACADÉMICA COCHABAMBA', size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('SEIS GUÍAS DEL\nPROCESO FORENSE DIGITAL')
r.font.size = Pt(26); r.bold = True; r.font.color.rgb = NAVY; r.font.name = BODY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('RFC 3227  ·  SWGDE  ·  ACPO  ·  CP4DF  ·  NIST SP 800-86  ·  ISO/IEC 27037')
r.font.size = Pt(12); r.font.color.rgb = ACCENT; r.font.name = BODY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(56)
r = p.add_run('Autoría, alcance, procesos, etapas, contenido de cada etapa\ny ejemplo de aplicación')
r.font.size = Pt(10.5); r.italic = True; r.font.color.rgb = GREY; r.font.name = BODY

t = doc.add_table(rows=0, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
for k, v in [('Materia', 'Informática Forense'),
             ('Tema', 'Guías y estándares del proceso forense digital'),
             ('Docente', '________________________________'),
             ('Estudiante', '________________________________'),
             ('Carrera', 'Ingeniería de Sistemas'),
             ('Lugar y fecha', 'Cochabamba – Bolivia, 2026')]:
    cells = t.add_row().cells
    cell_text(cells[0], k.upper(), bold=True, size=9.5, color=NAVY)
    cell_text(cells[1], v, size=10.5)
    cells[0].width = Cm(4.0); cells[1].width = Cm(11.9)
fixed_layout(t, [4.0, 11.9])
page_break()

# ══════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Índice', level=1)
for it in [
    '1. Resumen', '2. Introducción', '3. Objetivos',
    '4. Las seis guías de un vistazo',
    '5. ¿Quiénes desarrollan cada guía?',
    '     5.1 RFC 3227      5.2 SWGDE      5.3 ACPO',
    '     5.4 CP4DF      5.5 NIST SP 800-86      5.6 ISO/IEC 27037',
    '6. Alcance de cada guía',
    '     6.1 a 6.6  Alcance, destinatarios y límites de cada documento',
    '7. Procesos, etapas y contenido de cada etapa',
    '     7.1 RFC 3227: cinco etapas de recolección y archivo',
    '     7.2 SWGDE: cinco etapas del procedimiento de laboratorio',
    '     7.3 ACPO: cuatro principios y cuatro etapas operativas',
    '     7.4 CP4DF: tres fases del modelo de proceso común',
    '     7.5 NIST SP 800-86: cuatro fases del proceso forense',
    '     7.6 ISO/IEC 27037: cuatro procesos y dos roles',
    '8. Mapeo comparativo de las etapas',
    '9. Ejemplo de aplicación: fuga de información',
    '10. Recomendaciones', '11. Conclusiones', '12. Bibliografía',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(it); r.font.name = BODY; r.font.size = Pt(11)
    if not it.startswith('  '): r.bold = True
page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. RESUMEN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Resumen', level=1)
para('El presente trabajo analiza las seis guías que estructuran el trabajo de la informática forense: '
     'el RFC 3227 de la IETF, las mejores prácticas del SWGDE, la guía de buenas prácticas del ACPO '
     'británico, el modelo de proceso común CP4DF, la publicación especial NIST SP 800-86 y la norma '
     'internacional ISO/IEC 27037.')
para('De cada una se identifica quién la desarrolla, cuál es su alcance, qué proceso y etapas establece '
     'y qué contiene concretamente cada etapa. El análisis muestra que las seis no son alternativas '
     'excluyentes sino referencias complementarias: unas indican qué hacer en la escena, otras regulan '
     'quién puede hacerlo y con qué calidad, y otras ordenan el proceso completo de principio a fin.')
para('Para demostrar esa complementariedad, el trabajo desarrolla un caso hipotético de fuga de '
     'información en el que las seis guías intervienen en momentos distintos de una misma '
     'investigación, desde la llegada a la escena hasta la presentación del informe. Se concluye que '
     'la pregunta relevante no es cuál guía adoptar, sino en qué tramo del caso aplica cada una.')
para('Palabras clave: informática forense, evidencia digital, cadena de custodia, orden de volatilidad, '
     'guías y estándares forenses.', italic=True, size=10)

# ══════════════════════════════════════════════════════════════════════════
# 2. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Introducción', level=1)
para('La informática forense identifica, preserva, analiza y presenta evidencia digital de manera que '
     'resulte válida en un proceso legal o en una investigación interna. Su dificultad principal no '
     'está en las herramientas —hoy abundantes y en buena parte libres— sino en el método: una '
     'evidencia mal adquirida, mal documentada o analizada sin un procedimiento reproducible pierde su '
     'valor probatorio por más sólido que parezca el hallazgo técnico.')
para('Por esa razón la disciplina se apoya en guías y estándares que normalizan el trabajo del '
     'investigador. Seis de ellos concentran la mayor parte de las referencias del área y provienen de '
     'orígenes muy distintos: la comunidad técnica de internet, los laboratorios forenses '
     'estadounidenses, la policía británica, la academia alemana, la agencia de estándares '
     'norteamericana y el organismo internacional de normalización. Ese origen explica el enfoque de '
     'cada documento y, sobre todo, explica por qué ninguno cubre por sí solo todo el problema.')
para('Este trabajo los estudia en conjunto. La estructura sigue las seis preguntas planteadas: quiénes '
     'los desarrollan, cuál es su alcance, qué procesos y etapas establecen, qué contiene cada etapa, '
     'cómo se aplican sobre un caso concreto y qué conclusiones pueden extraerse de la comparación.')

# ══════════════════════════════════════════════════════════════════════════
# 3. OBJETIVOS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Objetivos', level=1)
doc.add_heading('3.1 Objetivo general', level=2)
para('Analizar de forma comparada las seis guías del proceso forense digital —RFC 3227, SWGDE, ACPO, '
     'CP4DF, NIST SP 800-86 e ISO/IEC 27037— identificando su autoría, alcance, procesos y etapas, el '
     'contenido de cada etapa y su aplicación práctica sobre un caso de incidente de seguridad.')
doc.add_heading('3.2 Objetivos específicos', level=2)
bullets([
    'Identificar el organismo, los autores, el país y el año de publicación de cada una de las seis guías.',
    'Determinar el alcance de cada documento: a quién se dirige, qué cubre y qué deja explícitamente fuera.',
    'Describir el proceso y las etapas que cada guía establece, y el contenido concreto de cada etapa.',
    'Construir un mapeo comparativo que alinee las etapas equivalentes de las seis guías.',
    'Aplicar las seis guías sobre un mismo caso de investigación forense y evidenciar en qué momento '
    'interviene cada una.',
    'Extraer conclusiones sobre su complementariedad y recomendaciones para su uso combinado.',
])

# ══════════════════════════════════════════════════════════════════════════
# 4. PANORAMA
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Las seis guías de un vistazo', level=1)
para('Antes del análisis detallado conviene fijar el panorama general. La tabla siguiente resume el '
     'origen, el enfoque y el año de cada documento.')
make_table(
    ['Guía', 'Origen', 'Enfoque principal', 'Año'],
    [['RFC 3227', 'IETF', 'Cómo recolectar y archivar la evidencia; define el orden de volatilidad', '2002'],
     ['SWGDE', 'Estados Unidos', 'Calidad y estandarización del laboratorio forense digital', '1998 en adelante'],
     ['ACPO', 'Reino Unido', 'Cuatro principios de conducta para quien manipula la evidencia', '2012 (v5)'],
     ['CP4DF', 'Académico (Alemania)', 'Proceso común que une la respuesta a incidentes con la pericia forense', '2007'],
     ['NIST SP 800-86', 'NIST – Estados Unidos', 'Proceso de cuatro fases integrado a la respuesta a incidentes', '2006'],
     ['ISO/IEC 27037', 'Internacional', 'Identificar, recolectar, adquirir y preservar la evidencia digital', '2012']],
    [3.0, 3.3, 7.2, 2.4], bold_first=True)
caption('Tabla 1. Panorama general de las seis guías analizadas.')
para('Una observación previa que atraviesa todo el trabajo: las guías no se ordenan por jerarquía sino '
     'por el tramo del proceso que cubren. Cinco de las seis son de descarga libre y gratuita; '
     'únicamente la norma ISO/IEC 27037 se adquiere mediante compra.')

# ══════════════════════════════════════════════════════════════════════════
# 5. QUIÉNES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. ¿Quiénes desarrollan cada guía?', level=1)
para('El origen institucional de cada documento condiciona su contenido. Una guía escrita por '
     'ingenieros de internet, una escrita por directores de laboratorios criminalísticos y una escrita '
     'por un organismo de normalización responden a preocupaciones distintas, y eso se refleja en lo '
     'que cada una decide detallar.')

doc.add_heading('5.1 RFC 3227 — IETF', level=2)
para('Fue redactado por Dominique Brezinski y Tom Killalea dentro del Network Working Group de la '
     'Internet Engineering Task Force (IETF), el organismo que define los estándares técnicos de '
     'internet. Se publicó en febrero de 2002 y se clasificó como BCP 55, es decir, Best Current '
     'Practice número 55, la categoría que la IETF reserva para las prácticas recomendadas y no para '
     'las especificaciones de protocolo.')

doc.add_heading('5.2 SWGDE — Scientific Working Group on Digital Evidence', level=2)
para('El SWGDE fue creado en 1998 por el grupo de directores de laboratorios criminalísticos federales '
     'de los Estados Unidos. Reúne a organizaciones dedicadas al análisis de evidencia digital y '
     'multimedia, entre ellas el FBI, la DEA, el Servicio Secreto y el propio NIST. Trabaja de forma '
     'coordinada con la IOCE (International Organization on Computer Evidence) para armonizar criterios '
     'a nivel internacional. No produce un documento único sino un cuerpo de publicaciones que se '
     'revisa y amplía de forma continua.')

doc.add_heading('5.3 ACPO — Association of Chief Police Officers', level=2)
para('La Good Practice Guide for Digital Evidence fue elaborada por la Association of Chief Police '
     'Officers del Reino Unido, con apoyo técnico de la empresa especializada 7Safe. Su versión 5, la '
     'más citada, se publicó en 2012. La ACPO se disolvió en 2015 y sus funciones pasaron al National '
     'Police Chiefs’ Council (NPCC), pero la guía mantiene su vigencia y se sigue citando con el '
     'nombre original.')

doc.add_heading('5.4 CP4DF — Modelo de proceso común', level=2)
para('El modelo fue propuesto por Felix C. Freiling y Bastian Schwittay, de la Universidad de Mannheim '
     '(Alemania), y presentado en 2007 en la conferencia IT Incident Management and IT Forensics (IMF). '
     'A diferencia de las demás referencias, se trata de una propuesta académica publicada en un '
     'congreso científico y no de un documento normativo emitido por un organismo oficial. Su valor '
     'está en la solidez conceptual del modelo, no en una autoridad institucional que lo respalde.')

doc.add_heading('5.5 NIST SP 800-86 — National Institute of Standards and Technology', level=2)
para('Fue elaborado por Karen Kent, Suzanne Chevalier, Tim Grance y Hung Dang, investigadores de la '
     'Computer Security Division perteneciente al Information Technology Laboratory del NIST, agencia '
     'federal dependiente del Departamento de Comercio de los Estados Unidos. Se publicó en 2006 dentro '
     'de la serie Special Publications 800, dedicada a la seguridad de la información.')

doc.add_heading('5.6 ISO/IEC 27037 — ISO e IEC', level=2)
para('La norma fue desarrollada por el subcomité conjunto ISO/IEC JTC 1/SC 27, dedicado a las técnicas '
     'de seguridad de la información, a través de su grupo de trabajo 4, con la participación de '
     'expertos designados por los países miembros. Se publicó en octubre de 2012. En Bolivia, el '
     'organismo nacional de normalización que representa a la ISO es IBNORCA, a través del cual se '
     'accede formalmente a la norma.')

# ══════════════════════════════════════════════════════════════════════════
# 6. ALCANCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Alcance de cada guía', level=1)
para('El alcance responde a dos preguntas: a quién se dirige el documento y qué tramo del problema '
     'cubre. Tan importante como lo que cada guía abarca es lo que deja deliberadamente fuera, porque '
     'ahí se ve por qué es necesario combinarlas.')

for h, dest, cubre, no in [
    ('6.1 RFC 3227',
     'Administradores de sistemas, personal de seguridad y equipos de respuesta a incidentes; en general, quien se encuentra con un sistema comprometido y debe actuar de inmediato.',
     'La recolección y el archivo de la evidencia durante un incidente: qué capturar, en qué orden, con qué precauciones y cómo conservarlo después.',
     'No cubre el análisis de la evidencia, no describe herramientas concretas y no constituye asesoría legal para ninguna jurisdicción en particular.'),
    ('6.2 SWGDE',
     'Laboratorios forenses digitales, sus examinadores y sus responsables de calidad; también instituciones que buscan acreditar su capacidad forense.',
     'La estandarización y la calidad del trabajo del laboratorio: procedimientos operativos escritos, validación de herramientas, competencia del personal y revisión de informes.',
     'No define un proceso de investigación único ni resuelve cuestiones de admisibilidad legal; se concentra en cómo debe operar el laboratorio que produce el resultado.'),
    ('6.3 ACPO',
     'Personal policial y agentes que manipulan evidencia digital en el Reino Unido, desde el primer respondiente en la escena hasta el perito que declara en el tribunal.',
     'La conducta correcta frente a la evidencia digital y las etapas del trabajo policial: planificación, captura en la escena, análisis y presentación ante el tribunal.',
     'Está redactada para el marco legal británico, por lo que sus referencias procesales no aplican directamente en otros países; sus principios, en cambio, sí son universales.'),
    ('6.4 CP4DF',
     'Equipos que deben atender un incidente y, al mismo tiempo, producir evidencia utilizable; también investigadores y docentes del área.',
     'Un modelo conceptual que integra en un solo proceso la respuesta a incidentes y la pericia forense, con un análisis organizado como ciclo de hipótesis y verificación.',
     'Al ser un modelo académico, no prescribe herramientas, formatos ni procedimientos operativos detallados: aporta la estructura de razonamiento, no el manual de campo.'),
    ('6.5 NIST SP 800-86',
     'Equipos de respuesta a incidentes, investigadores forenses, administradores de sistemas y organizaciones que necesitan definir políticas forenses propias.',
     'La integración de las técnicas forenses en la respuesta a incidentes y la construcción de una capacidad forense organizacional, con un proceso de cuatro fases.',
     'No es asesoría legal, no reemplaza la normativa procesal de cada país y no es el manual de ninguna herramienta en particular.'),
    ('6.6 ISO/IEC 27037',
     'Primeros respondientes y especialistas en evidencia digital, así como organizaciones que buscan alinearse con una norma internacional certificable.',
     'Únicamente las etapas iniciales del tratamiento de la evidencia: identificación, recolección, adquisición y preservación, junto con los roles que las ejecutan.',
     'No cubre el análisis ni la presentación de resultados: esas etapas corresponden a las normas ISO/IEC 27042 y 27043, y la gestión del incidente a la 27035.'),
]:
    doc.add_heading(h, level=2)
    bullets([('Destinatarios.', dest), ('Qué cubre.', cubre), ('Qué deja fuera.', no)])

# ══════════════════════════════════════════════════════════════════════════
# 7. PROCESOS Y ETAPAS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Procesos, etapas y contenido de cada etapa', level=1)
para('Esta sección responde de forma conjunta las preguntas tres y cuatro. Para cada guía se presenta '
     'la secuencia de etapas que establece y, dentro de cada etapa, el contenido concreto: las '
     'actividades que efectivamente se ejecutan.')

# 7.1 RFC 3227
doc.add_heading('7.1 RFC 3227: cinco etapas de recolección y archivo', level=2)
para('El RFC 3227 es el documento más breve y más operativo del conjunto. No propone un ciclo de '
     'investigación completo sino un procedimiento para los primeros minutos, que es donde se pierde '
     'la mayor parte de la evidencia.')
make_table(['Etapa', 'Contenido'],
    [['1. Principios rectores',
      'Capturar una imagen lo más fiel posible del sistema; anotar todo con fecha y hora indicando si es local o UTC; minimizar los cambios en los datos; recolectar primero y analizar después; proceder de lo más volátil a lo menos volátil; y estar preparado para declarar ante un tribunal.'],
     ['2. Orden de volatilidad',
      'Secuencia obligatoria de captura: registros y caché; tabla de rutas, caché ARP, tabla de procesos, estadísticas del núcleo y memoria; sistemas de archivos temporales; disco; registros y monitoreo remotos; configuración física y topología de la red; y por último los medios de respaldo.'],
     ['3. Procedimiento de recolección',
      'El método debe ser transparente y verificable. Se usan herramientas propias almacenadas en un medio de solo lectura, sin confiar en los binarios del sistema comprometido. Se prohíbe apagar el equipo antes de recolectar, ejecutar programas que alteren las marcas de acceso y cerrar procesos en curso.'],
     ['4. Procedimiento de archivado',
      'La cadena de custodia debe registrar dónde, cuándo y quién descubrió y recolectó la evidencia; quién la manejó y por cuánto tiempo; cómo se almacenó; y cuándo y cómo se produjo cada transferencia. El material se conserva en un medio seguro con acceso restringido y registrado.'],
     ['5. Marco legal y de privacidad',
      'La evidencia debe cumplir cinco atributos: ser admisible, auténtica, completa, confiable y creíble. No debe invadirse la privacidad de las personas sin justificación ni autorización previa, y toda recolección debe respetar la normativa aplicable.']],
    [3.6, 12.3])
caption('Tabla 2. Etapas del RFC 3227 y su contenido.')

# 7.2 SWGDE
doc.add_heading('7.2 SWGDE: cinco etapas del procedimiento de laboratorio', level=2)
para('El SWGDE no publica un proceso único, pero de sus documentos de mejores prácticas y de su '
     'procedimiento operativo modelo se desprende una secuencia de trabajo estable, centrada en el '
     'laboratorio y en la calidad del resultado.')
make_table(['Etapa', 'Contenido'],
    [['1. Preparación del laboratorio',
      'Procedimientos operativos estándar escritos, aprobados por la dirección y revisados periódicamente; validación documentada de cada herramienta de hardware y software antes de emplearla en casos reales; y personal con capacitación y competencia técnica demostrada.'],
     ['2. Inspección visual y documentación',
      'Al recibir la evidencia se registra marca, modelo, número de serie, estado físico y daños visibles del dispositivo. Se fotografía y se etiqueta antes de cualquier manipulación técnica, y se abre el expediente del caso.'],
     ['3. Duplicación forense',
      'Se adquiere la imagen del medio empleando bloqueo de escritura, se calcula el valor hash del original y de la copia, y se verifica su coincidencia. Solo después de esa verificación se autoriza el trabajo sobre la copia.'],
     ['4. Examen del medio',
      'Todo el análisis se ejecuta sobre la copia verificada. Cada acción realizada queda asentada en la bitácora del caso, indicando la herramienta empleada y su versión exacta, de modo que el trabajo sea reproducible por otro examinador.'],
     ['5. Devolución y control de calidad',
      'Retorno documentado de la evidencia a su custodio, revisión técnica y administrativa del informe por un segundo examinador independiente, y archivo completo del expediente con todos sus registros.']],
    [3.6, 12.3])
caption('Tabla 3. Etapas del procedimiento SWGDE y su contenido.')
para('El rasgo distintivo del SWGDE es la etapa cinco: la revisión por un segundo examinador. Ninguna '
     'otra de las seis guías incorpora de forma explícita un control cruzado del informe antes de su '
     'emisión.')

# 7.3 ACPO
doc.add_heading('7.3 ACPO: cuatro principios y cuatro etapas operativas', level=2)
para('El ACPO es la única guía del conjunto que comienza regulando la conducta y no el procedimiento. '
     'Sus cuatro principios son su aporte central y condicionan todo lo demás.')
make_table(['Principio', 'Contenido'],
    [['P1. No alterar los datos',
      'Ninguna acción de la autoridad, de las personas empleadas por ella o de sus agentes debe modificar los datos que posteriormente puedan presentarse ante un tribunal.'],
     ['P2. Competencia justificada',
      'Cuando resulte imprescindible acceder a los datos originales, quien lo haga debe ser competente para ello y capaz de declarar explicando la relevancia y las implicaciones de sus acciones.'],
     ['P3. Registro de auditoría',
      'Debe crearse y preservarse un registro de todos los procesos aplicados a la evidencia digital. Un tercero independiente debe poder examinar esos procesos, repetirlos y obtener el mismo resultado.'],
     ['P4. Responsabilidad del jefe de caso',
      'La persona a cargo de la investigación tiene la responsabilidad general de garantizar que la ley y estos principios se cumplan por parte de todo el equipo interviniente.']],
    [4.2, 11.7])
caption('Tabla 4. Los cuatro principios del ACPO.')
para('Sobre esos principios, la guía organiza el trabajo operativo en cuatro etapas:')
make_table(['Etapa', 'Contenido'],
    [['1. Planificación',
      'Obtención de la autorización legal correspondiente, preparación del operativo, asignación de roles y briefing del equipo que intervendrá en la escena.'],
     ['2. Captura en la escena',
      'Fotografiado y etiquetado de los dispositivos, decisión entre incautar el equipo o realizar una captura en vivo cuando apagarlo destruiría evidencia, y apertura del registro de auditoría.'],
     ['3. Análisis',
      'Examen sobre la copia por parte de personal competente, con registro de cada proceso aplicado de modo que un tercero pueda repetirlo y obtener el mismo resultado.'],
     ['4. Presentación',
      'Elaboración del informe y de la declaración testimonial, con el detalle necesario para que el tribunal comprenda el alcance y las limitaciones de los hallazgos.']],
    [3.6, 12.3])
caption('Tabla 5. Etapas operativas del ACPO y su contenido.')
para('La continuidad de la evidencia y el registro de auditoría no constituyen una etapa: atraviesan '
     'las cuatro y deben mantenerse sin interrupción desde el primer contacto hasta la declaración '
     'final.')

# 7.4 CP4DF
doc.add_heading('7.4 CP4DF: tres fases del modelo de proceso común', level=2)
para('El modelo parte de una constatación práctica: cuando ocurre un incidente hay dos equipos con '
     'objetivos que chocan. Uno quiere restablecer el servicio cuanto antes; el otro quiere preservar '
     'la evidencia intacta. El CP4DF los reúne en un solo proceso de tres fases.')
make_table(['Fase', 'Contenido'],
    [['1. Pre-análisis (Pre-Analysis)',
      'Preparación previa al incidente: políticas, herramientas y capacitación. Detección del incidente y respuesta inicial. Formulación de la estrategia de investigación según el tipo de caso. Recolección y preservación de los datos volátiles antes de que se pierdan.'],
     ['2. Análisis (Analysis)',
      'Examen de los datos recolectados, formulación de una hipótesis sobre lo ocurrido y verificación o refutación de esa hipótesis contra la evidencia disponible. El ciclo se repite: si la hipótesis se refuta, se vuelve al examen con un enfoque distinto, hasta que una resiste el contraste.'],
     ['3. Post-análisis (Post-Analysis)',
      'Documentación completa del procedimiento seguido, presentación de los resultados a los destinatarios correspondientes, restauración del servicio afectado y registro de las lecciones aprendidas para mejorar la respuesta futura.']],
    [4.6, 11.3])
caption('Tabla 6. Fases del modelo CP4DF y su contenido.')
para('El aporte propio del modelo es doble. Por un lado reconcilia la urgencia operativa con el rigor '
     'probatorio, evitando que la contención del incidente destruya la evidencia. Por otro convierte el '
     'análisis en un ciclo de hipótesis y verificación —esto es, en aplicación del método científico— '
     'en lugar de una búsqueda lineal de indicios.')

# 7.5 NIST
doc.add_heading('7.5 NIST SP 800-86: cuatro fases del proceso forense', level=2)
para('El NIST propone el proceso más difundido del conjunto. Su valor no está solo en dividir el '
     'trabajo en etapas, sino en definir qué recibe y qué entrega cada una, mediante una cadena de '
     'transformación progresiva del dato.')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
r = p.add_run('MEDIOS  →  DATOS  →  INFORMACIÓN  →  EVIDENCIA')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT; r.font.name = BODY
make_table(['Fase', 'Contenido'],
    [['1. Recolección (Collection)',
      'Identificar, etiquetar, registrar y adquirir los datos de las fuentes relevantes, preservando su integridad. Incluye la imagen bit a bit con bloqueador de escritura, la verificación mediante hash y la apertura de la cadena de custodia.'],
     ['2. Examen (Examination)',
      'Procesar los datos adquiridos para hacerlos manejables: descartar archivos conocidos del sistema mediante listas de referencia, recuperar archivos borrados y espacio no asignado, y filtrar por palabras clave, fechas y tipos de archivo.'],
     ['3. Análisis (Analysis)',
      'Correlacionar la información obtenida y derivar conclusiones: construir la línea de tiempo del incidente, cruzar registros de red, sistema operativo y aplicaciones, e identificar artefactos como malware, persistencia y datos exfiltrados.'],
     ['4. Reporte (Reporting)',
      'Documentar el proceso y los hallazgos en tres productos: informe técnico reproducible, informe ejecutivo sin lenguaje técnico y anexo legal con la cadena de custodia completa y los hashes de la evidencia.']],
    [4.6, 11.3])
caption('Tabla 7. Fases del NIST SP 800-86 y su contenido.')

# 7.6 ISO
doc.add_heading('7.6 ISO/IEC 27037: cuatro procesos y dos roles', level=2)
para('La norma internacional se concentra en las etapas iniciales, que identifica con las siglas ICAP: '
     'identificación, recolección, adquisición y preservación.')
make_table(['Proceso', 'Contenido'],
    [['1. Identificación',
      'Localizar y reconocer los dispositivos y medios que pueden contener evidencia, tanto en su forma física como lógica, y priorizar su tratamiento según el grado de volatilidad de los datos que albergan.'],
     ['2. Recolección',
      'Retirar los dispositivos del lugar de los hechos y trasladarlos a un entorno controlado, cuando la decisión adoptada es llevarse el original en lugar de copiarlo en el sitio.'],
     ['3. Adquisición',
      'Producir una copia verificable de la evidencia digital y documentar el método empleado, cuando no es viable o no es conveniente llevarse el original, por ejemplo en servidores que no pueden detenerse.'],
     ['4. Preservación',
      'Proteger la integridad y la originalidad de la evidencia durante el embalaje, el traslado y el almacenamiento, mediante materiales adecuados, etiquetado y una cadena de custodia sin interrupciones.']],
    [3.8, 12.1])
caption('Tabla 8. Los cuatro procesos de la ISO/IEC 27037 y su contenido.')
para('La norma define además dos roles con responsabilidades diferenciadas, algo que ninguna otra guía '
     'formaliza de este modo:')
bullets([
    ('DEFR — Digital Evidence First Responder.', 'Primer respondiente de evidencia digital: persona '
     'autorizada, capacitada y calificada para actuar primero en la escena, recolectando y adquiriendo '
     'la evidencia bajo su responsabilidad.'),
    ('DES — Digital Evidence Specialist.', 'Especialista en evidencia digital: profesional con '
     'conocimientos y habilidades ampliadas para resolver situaciones técnicamente complejas, como '
     'arreglos RAID, adquisición en redes, servidores de correo o sistemas propietarios.'),
])
para('La ISO/IEC 27037 exige además que todo el tratamiento satisfaga los principios de relevancia, '
     'confiabilidad, suficiencia, auditabilidad, repetibilidad, reproducibilidad y justificabilidad, y '
     'que se minimice en todo momento el manejo de la evidencia digital original.')

# ══════════════════════════════════════════════════════════════════════════
# 8. MAPEO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Mapeo comparativo de las etapas', level=1)
para('Las seis guías emplean nombres distintos para actividades equivalentes. La tabla siguiente alinea '
     'sus etapas sobre cinco momentos comunes del trabajo forense, y muestra con un guion los tramos '
     'que cada documento deliberadamente no cubre.')
make_table(
    ['Guía', 'Preparación', 'Identificación y recolección', 'Adquisición y preservación', 'Examen y análisis', 'Informe y presentación'],
    [['RFC 3227', 'Kit de herramientas', 'Orden de volatilidad', 'Archivado y custodia', '—', 'Notas para declarar'],
     ['SWGDE', 'SOP y validación', 'Inspección visual', 'Duplicación forense', 'Examen del medio', 'Informe revisado'],
     ['ACPO', 'Planificación', 'Captura en escena', 'Continuidad de la evidencia', 'Personal competente', 'Declaración'],
     ['CP4DF', 'Pre-análisis', 'Pre-análisis', 'Pre-análisis', 'Análisis iterativo', 'Post-análisis'],
     ['NIST SP 800-86', 'Capacidad forense', 'Collection', 'Collection', 'Examination y Analysis', 'Reporting'],
     ['ISO/IEC 27037', '—', 'Identificación y recolección', 'Adquisición y preservación', 'Lo cubre la 27042', 'Lo cubre la 27043']],
    [2.8, 2.6, 2.7, 2.7, 2.6, 2.5], size=8, bold_first=True)
caption('Tabla 9. Alineación de las etapas equivalentes de las seis guías.')
para('La lectura del mapeo permite tres observaciones. Primero, solo el SWGDE, el ACPO, el CP4DF y el '
     'NIST cubren el ciclo completo. Segundo, el RFC 3227 y la ISO/IEC 27037 son deliberadamente '
     'parciales: se concentran en la escena y en la preservación, que es donde el error resulta '
     'irreversible. Tercero, allí donde una guía se detiene, otra continúa, lo que confirma que están '
     'pensadas para usarse en conjunto.')

# ══════════════════════════════════════════════════════════════════════════
# 9. EJEMPLO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Ejemplo de aplicación: fuga de información', level=1)
doc.add_heading('9.1 Planteamiento del caso', level=2)
para('La gerencia de una empresa mediana denuncia que el listado completo de sus clientes apareció en '
     'poder de la competencia. La sospecha recae sobre un funcionario del área comercial. Se solicita '
     'una investigación forense que determine si hubo salida de información, por qué medio y en qué '
     'momento, con material que pueda sostenerse ante la autoridad competente.')
para('Las condiciones de la escena presentan tres complicaciones concretas. El equipo del sospechoso se '
     'encuentra encendido y conectado a la red corporativa, de modo que apagarlo destruiría la '
     'evidencia en memoria. Existe además un servidor de archivos con arreglo RAID que no puede '
     'trasladarse sin detener la operación de la empresa. Por último, los registros del servidor proxy '
     'rotan cada setenta y dos horas, por lo que el tiempo apremia.')

doc.add_heading('9.2 Aplicación de las seis guías sobre el mismo caso', level=2)
para('La investigación no elige una guía: las aplica todas, cada una en el momento que le corresponde. '
     'La tabla siguiente sigue el caso en orden cronológico.')
make_table(
    ['Momento', 'Acción ejecutada', 'Guía que la rige'],
    [['Llegada a la escena', 'Se fotografía el puesto de trabajo, se etiquetan los equipos y se abre el registro de auditoría del caso', 'ACPO (P3) y SWGDE'],
     ['Equipo encendido', 'Se vuelca la memoria RAM y se registran las conexiones activas antes de tocar el disco', 'RFC 3227'],
     ['Acceso al equipo vivo', 'Interviene únicamente personal competente, que documenta y justifica cada comando ejecutado', 'ACPO (P2)'],
     ['Decisión sobre el original', 'El DEFR asegura la escena y convoca al DES para resolver el arreglo RAID que no puede trasladarse', 'ISO/IEC 27037'],
     ['Adquisición', 'Imagen bit a bit con bloqueador de escritura y verificación mediante hash SHA-256 del original y la copia', 'ISO/IEC 27037 y SWGDE'],
     ['Traslado y resguardo', 'Embalaje adecuado, etiquetado y firma del acta de cadena de custodia en cada transferencia', 'RFC 3227 e ISO/IEC 27037'],
     ['Organización del trabajo', 'El caso se estructura en las fases de recolección, examen, análisis y reporte, con entregables definidos', 'NIST SP 800-86'],
     ['Examen', 'Descarte de archivos conocidos, recuperación de borrados y filtrado por rango de fechas y palabras clave', 'NIST SP 800-86 y SWGDE'],
     ['Análisis', 'Se formula la hipótesis de la salida por correo personal y se verifica contra los registros del proxy', 'CP4DF y NIST SP 800-86'],
     ['Revisión del informe', 'Un segundo examinador realiza la revisión técnica y administrativa antes de la emisión', 'SWGDE'],
     ['Presentación', 'Informe técnico reproducible, informe ejecutivo y declaración testimonial ante la autoridad', 'ACPO y NIST SP 800-86']],
    [3.4, 8.5, 4.0], size=9)
caption('Tabla 10. Recorrido cronológico del caso y guía que rige cada acción.')

doc.add_heading('9.3 Resultado', level=2)
para('El análisis confirma la hipótesis inicial: la información salió a través de una cuenta de correo '
     'personal, en tres envíos realizados fuera del horario laboral, y los registros del proxy '
     'corroboran de forma independiente lo hallado en el equipo. La primera hipótesis considerada —una '
     'copia a dispositivo USB— se descartó al no encontrar rastros de montaje en el registro del '
     'sistema, y el ciclo volvió al examen con un enfoque distinto, tal como prevé el modelo CP4DF.')
para('El informe resultante resiste tanto el cuestionamiento técnico como el legal, y la razón es '
     'precisamente que ninguna decisión quedó librada al criterio del momento: cada una está respaldada '
     'por una guía reconocida. El orden de captura lo fija el RFC 3227; la conducta del personal, el '
     'ACPO; la adquisición y la preservación, la ISO/IEC 27037; la calidad del laboratorio y la '
     'revisión cruzada, el SWGDE; la estructura del trabajo, el NIST; y el razonamiento del análisis, '
     'el CP4DF.')

# ══════════════════════════════════════════════════════════════════════════
# 10. RECOMENDACIONES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Recomendaciones', level=1)
bullets([
    ('Adoptar una guía marco y complementarla, no elegir una sola.', 'Conviene tomar el NIST SP 800-86 '
     'como estructura general del proceso y apoyarse en las demás para los tramos que este no detalla.'),
    ('Preparar la capacidad forense antes del incidente.', 'Políticas, roles, herramientas validadas y '
     'personal capacitado deben estar resueltos con anticipación. Improvisar durante el incidente es la '
     'causa más frecuente de pérdida de evidencia.'),
    ('Respetar siempre el orden de volatilidad del RFC 3227.', 'Capturar la memoria y el estado de la '
     'red antes de apagar el equipo. Apagar primero y pensar después destruye evidencia irrecuperable.'),
    ('Aplicar los cuatro principios del ACPO como regla de conducta permanente.', 'No alterar, '
     'justificar el acceso, dejar rastro auditable y asignar responsabilidad, con independencia de la '
     'jurisdicción en la que se trabaje.'),
    ('Formalizar los roles DEFR y DES de la ISO/IEC 27037.', 'Definir por escrito quién puede actuar '
     'en la escena y a quién se convoca ante una situación técnicamente compleja evita improvisaciones '
     'costosas.'),
    ('Incorporar la revisión cruzada del SWGDE.', 'Que un segundo examinador revise el informe antes '
     'de su emisión es la práctica de control de calidad más barata y más eficaz del conjunto.'),
    ('Organizar el análisis como ciclo de hipótesis y verificación, según el CP4DF.', 'Escribir la '
     'hipótesis antes de buscar la evidencia reduce el sesgo de confirmación y hace explícito el '
     'razonamiento en el informe.'),
    ('Documentar en el momento, no al final.', 'Quién, cuándo, con qué herramienta y con qué '
     'parámetros. Lo que no está documentado, a efectos prácticos, no ocurrió.'),
])

# ══════════════════════════════════════════════════════════════════════════
# 11. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Conclusiones', level=1)
bullets([
    'Las seis guías analizadas no compiten entre sí: cubren tramos distintos de un mismo problema y en '
    'la práctica profesional se aplican de forma combinada, no excluyente.',
    'El RFC 3227 y la ISO/IEC 27037 son las más operativas en la escena. El primero aporta el orden de '
    'volatilidad, que es la regla que evita destruir evidencia en los primeros minutos; la segunda '
    'formaliza los roles del interviniente y los requisitos de preservación.',
    'El SWGDE y el ACPO regulan al actor antes que al procedimiento. Uno establece la calidad del '
    'laboratorio que produce el resultado y el otro la conducta de quien manipula la evidencia; ambos '
    'atienden un aspecto que las guías de proceso dan por supuesto.',
    'El NIST SP 800-86 y el CP4DF son los modelos de proceso. El primero organiza el trabajo en fases '
    'con entradas y salidas definidas mediante la cadena medios – datos – información – evidencia; el '
    'segundo lo convierte en un ciclo de hipótesis y verificación e integra la respuesta a incidentes '
    'con la pericia forense.',
    'El origen institucional de cada documento explica su enfoque: la IETF se ocupa del procedimiento '
    'técnico, los laboratorios de la calidad, la policía de la conducta, la academia del razonamiento, '
    'la agencia de estándares del proceso y el organismo de normalización de la preservación.',
    'El ejemplo desarrollado demuestra que las seis pueden convivir en una sola investigación sin '
    'contradicción alguna, y que el informe resultante gana solidez precisamente porque cada decisión '
    'tomada está respaldada por una referencia reconocida.',
    'La pregunta relevante, en consecuencia, no es cuál guía adoptar sino en qué tramo del caso aplica '
    'cada una. Un investigador que conoce solo una de ellas queda expuesto exactamente en el tramo que '
    'esa guía no cubre.',
], style='List Number')

# ══════════════════════════════════════════════════════════════════════════
# 12. BIBLIOGRAFÍA
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Bibliografía', level=1)
refs = [
    'Association of Chief Police Officers (2012). ACPO Good Practice Guide for Digital Evidence, '
    'versión 5. Reino Unido: ACPO.',
    'Brezinski, D. y Killalea, T. (2002). Guidelines for Evidence Collection and Archiving (RFC 3227, '
    'BCP 55). Internet Engineering Task Force.',
    'Cichonski, P., Millar, T., Grance, T. y Scarfone, K. (2012). Computer Security Incident Handling '
    'Guide (NIST Special Publication 800-61 Rev. 2). Gaithersburg: National Institute of Standards and '
    'Technology.',
    'Freiling, F. C. y Schwittay, B. (2007). A Common Process Model for Incident Response and Computer '
    'Forensics. En Proceedings of the IT Incident Management and IT Forensics (IMF 2007). Stuttgart, '
    'Alemania.',
    'International Organization for Standardization (2012). ISO/IEC 27037: Information technology – '
    'Security techniques – Guidelines for identification, collection, acquisition and preservation of '
    'digital evidence. Ginebra: ISO.',
    'International Organization for Standardization (2015). ISO/IEC 27042: Guidelines for the analysis '
    'and interpretation of digital evidence. Ginebra: ISO.',
    'Kent, K., Chevalier, S., Grance, T. y Dang, H. (2006). Guide to Integrating Forensic Techniques '
    'into Incident Response (NIST Special Publication 800-86). Gaithersburg: National Institute of '
    'Standards and Technology.',
    'Scientific Working Group on Digital Evidence (2018). SWGDE Best Practices for Computer Forensic '
    'Acquisitions. Estados Unidos: SWGDE.',
    'Scientific Working Group on Digital Evidence (2019). SWGDE Model Standard Operation Procedures for '
    'Computer Forensics. Estados Unidos: SWGDE.',
]
for r_ in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(r_); run.font.name = BODY; run.font.size = Pt(10.5)

doc.sections[0].different_first_page_header_footer = True
add_page_number_footer(doc.sections[0])
doc.save('Guias_Forenses_Documento.docx')
print('OK — Guias_Forenses_Documento.docx')
