# -*- coding: utf-8 -*-
"""Genera el documento académico sobre la guía NIST SP 800-86."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x14, 0x2A, 0x45)
ACCENT = RGBColor(0x0E, 0x6C, 0x82)
GREY   = RGBColor(0x5A, 0x6A, 0x78)
BODY   = 'Calibri'

doc = Document()

# ── Configuración de página y estilos base ───────────────────────────────
sec = doc.sections[0]
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = BODY
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), BODY)
pf = normal.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for name, size, color, before, after in [
    ('Heading 1', 16, NAVY,   18, 8),
    ('Heading 2', 13, ACCENT, 14, 6),
    ('Heading 3', 11.5, NAVY, 10, 4),
]:
    st = doc.styles[name]
    st.font.name = BODY
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    st.paragraph_format.keep_with_next = True

# ── Utilidades ────────────────────────────────────────────────────────────
def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

def cell_text(cell, text, bold=False, size=9.5, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = BODY; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color

def fixed_layout(t, widths):
    t.autofit = False
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(Cm(w).twips)))

def repeat_header(t):
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true')
    trPr.append(th)

def no_row_split(t):
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true')
        trPr.insert(0, cs)

def make_table(headers, rows, widths, highlight=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, '142A45')
        cell_text(c, h, bold=True, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        hl = highlight is not None and ri == highlight
        for i, v in enumerate(row):
            if hl: shade(cells[i], 'DDEEF3')
            cell_text(cells[i], v, bold=(hl and i == 0), size=9.5)
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Cm(widths[i])
    fixed_layout(t, widths)
    repeat_header(t)
    no_row_split(t)
    return t

def para(text, style=None, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = BODY; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullets(items, style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.name = BODY
            r2 = p.add_run(' ' + it[1]); r2.font.name = BODY
        else:
            r = p.add_run(it); r.font.name = BODY

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = BODY; r.font.size = Pt(8.5); r.italic = True
    r.font.color.rgb = GREY

def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p.add_run('NIST SP 800-86  ·  Informática Forense  ·  Página ')
    r0.font.size = Pt(8); r0.font.name = BODY; r0.font.color.rgb = GREY
    r = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r._r.addnext(fld)
    rr = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '16'); rpr.append(sz)
    rr.append(rpr)
    t = OxmlElement('w:t'); t.text = '1'; rr.append(t)
    fld.append(rr)

# ══════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════
para('ESCUELA MILITAR DE INGENIERÍA', size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para('«Mcal. Antonio José de Sucre»', size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=4)
para('UNIDAD ACADÉMICA COCHABAMBA', size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('NIST SP 800-86')
r.font.size = Pt(30); r.bold = True; r.font.color.rgb = NAVY; r.font.name = BODY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Guía para integrar técnicas forenses\nen la respuesta a incidentes')
r.font.size = Pt(15); r.font.color.rgb = ACCENT; r.font.name = BODY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)
r = p.add_run('Guide to Integrating Forensic Techniques into Incident Response')
r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GREY; r.font.name = BODY

t = doc.add_table(rows=0, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
for k, v in [('Materia', 'Informática Forense'),
             ('Tema', 'Guías y estándares del proceso forense digital'),
             ('Docente', '________________________________'),
             ('Estudiante', '________________________________'),
             ('Carrera', 'Ingeniería de Sistemas'),
             ('Lugar y fecha', 'Cochabamba – Bolivia, 2026')]:
    cells = t.add_row().cells
    cell_text(cells[0], k.upper(), bold=True, size=9.5, color=NAVY)
    cell_text(cells[1], v, size=10.5)
    cells[0].width = Cm(4.0); cells[1].width = Cm(11.9)
fixed_layout(t, [4.0, 11.9])
page_break()

# ══════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Índice', level=1)
indice = [
    '1. Resumen',
    '2. Introducción',
    '3. Objetivos',
    '     3.1 Objetivo general',
    '     3.2 Objetivos específicos',
    '4. El NIST y la serie SP 800',
    '5. ¿Quiénes elaboran la guía?',
    '6. Alcance de la guía',
    '7. Panorama de las guías del área forense',
    '8. El proceso forense según el NIST SP 800-86',
    '     8.1 Visión general del proceso',
    '     8.2 Fase 1: Recolección (Collection)',
    '     8.3 Fase 2: Examen (Examination)',
    '     8.4 Fase 3: Análisis (Analysis)',
    '     8.5 Fase 4: Reporte (Reporting)',
    '9. Consideraciones transversales al proceso',
    '10. Ejemplo de aplicación: incidente de ransomware',
    '11. Recomendaciones',
    '12. Conclusiones',
    '13. Bibliografía',
]
for it in indice:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(it); r.font.name = BODY; r.font.size = Pt(11)
    if not it.startswith('  '): r.bold = True
page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. RESUMEN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Resumen', level=1)
para('El presente trabajo analiza la publicación especial NIST SP 800-86, titulada Guide to Integrating '
     'Forensic Techniques into Incident Response, elaborada por el Instituto Nacional de Estándares y '
     'Tecnología de los Estados Unidos y publicada en 2006. Se estudia el organismo que la produce, su '
     'autoría, el alcance que declara y, sobre todo, el proceso forense de cuatro fases que propone: '
     'recolección, examen, análisis y reporte.')
para('El documento describe el contenido concreto de cada fase —las actividades que se ejecutan, las '
     'herramientas típicas y el resultado que entrega a la siguiente etapa— y explica la cadena de '
     'transformación del dato que articula todo el modelo: los medios se convierten en datos, los datos '
     'en información y la información en evidencia. Para demostrar su aplicabilidad, el proceso completo '
     'se ejecuta sobre un caso hipotético de infección por ransomware en una empresa, incluyendo la '
     'reconstrucción de la línea de tiempo del ataque y el artefacto que respalda cada hecho.')
para('Se concluye que el SP 800-86 destaca entre las guías del área por tres razones: convierte la '
     'informática forense en un proceso repetible y auditable, integra la preservación de la evidencia '
     'dentro de la respuesta al incidente en lugar de tratarla como una actividad posterior, y es de '
     'acceso libre y gratuito, lo que la hace especialmente aplicable en contextos académicos y en '
     'organizaciones con recursos limitados.')
para('Palabras clave: informática forense, evidencia digital, respuesta a incidentes, cadena de custodia, '
     'NIST SP 800-86.', italic=True, size=10)

# ══════════════════════════════════════════════════════════════════════════
# 2. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Introducción', level=1)
para('La informática forense es la disciplina que se ocupa de identificar, preservar, analizar y presentar '
     'evidencia digital de manera que resulte válida en un proceso legal o en una investigación interna. '
     'Su dificultad no está tanto en las herramientas —hoy abundantes y en buena parte libres— sino en el '
     'método: una evidencia mal adquirida, mal documentada o analizada sin un procedimiento reproducible '
     'pierde todo su valor probatorio, por más sólido que parezca el hallazgo técnico.')
para('Precisamente por eso existen guías y estándares que normalizan el trabajo del investigador. En el '
     'ámbito de la materia se identifican seis referencias principales: RFC 3227, SWGDE, ACPO, CP4DF, '
     'NIST e ISO. Cada una nace de un contexto distinto —la comunidad técnica de internet, los '
     'laboratorios forenses estadounidenses, la policía británica, la academia alemana, la agencia de '
     'estándares norteamericana y el organismo internacional de normalización— y, en consecuencia, cada '
     'una pone el acento en un aspecto diferente del problema.')
para('De ese conjunto, este trabajo se centra en el NIST SP 800-86. La razón de la elección es práctica: '
     'es la guía que expresa el trabajo forense como un proceso de fases claramente delimitadas, con '
     'entradas y salidas identificables, lo que permite no solo explicarla sino demostrarla sobre un caso '
     'concreto. A ello se suma que el documento es de descarga libre y gratuita, a diferencia de la norma '
     'ISO/IEC 27037, que es de pago.')

# ══════════════════════════════════════════════════════════════════════════
# 3. OBJETIVOS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Objetivos', level=1)
doc.add_heading('3.1 Objetivo general', level=2)
para('Analizar la guía NIST SP 800-86 para comprender el proceso forense de cuatro fases que propone, el '
     'contenido de cada etapa y su aplicación práctica en la investigación de incidentes de seguridad '
     'informática.')
doc.add_heading('3.2 Objetivos específicos', level=2)
bullets([
    'Identificar el organismo que elabora la guía, sus autores y el alcance que el documento declara.',
    'Ubicar el SP 800-86 dentro del panorama de guías forenses vigentes y justificar su elección.',
    'Describir las cuatro fases del proceso forense y el contenido concreto de cada una de ellas.',
    'Reconocer las consideraciones transversales al proceso: cadena de custodia, orden de volatilidad y '
    'requisitos de admisibilidad.',
    'Aplicar el proceso completo a un caso de incidente de seguridad y extraer recomendaciones prácticas.',
])

# ══════════════════════════════════════════════════════════════════════════
# 4. EL NIST Y LA SERIE SP 800
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. El NIST y la serie SP 800', level=1)
para('El National Institute of Standards and Technology (NIST) es una agencia federal dependiente del '
     'Departamento de Comercio de los Estados Unidos. Fue fundado en 1901 con el nombre de National '
     'Bureau of Standards, con el encargo original de unificar los patrones de medida del país, y adoptó '
     'su denominación actual en 1988, cuando su mandato se amplió hacia la promoción de la innovación y '
     'la competitividad tecnológica.')
para('Conviene aclarar un punto que suele confundirse: el NIST no es un organismo policial ni un ente '
     'regulador. No investiga delitos ni impone sanciones. Lo que produce son estándares y guías '
     'técnicas. Su influencia es de naturaleza distinta: sus publicaciones son obligatorias para las '
     'agencias federales estadounidenses y, por su calidad y su gratuidad, terminaron adoptándose de '
     'forma voluntaria como referencia en buena parte del mundo.')
para('Dentro de su producción, la serie Special Publications 800 agrupa los documentos dedicados a la '
     'seguridad de la información. Reúne más de doscientos títulos de descarga libre en el portal '
     'csrc.nist.gov, entre los que se encuentran referencias muy conocidas como el SP 800-53 sobre '
     'controles de seguridad, el SP 800-61 sobre manejo de incidentes y el SP 800-86 que aquí se analiza. '
     'Estas dos últimas publicaciones son complementarias y conviene leerlas en conjunto: la 61 explica '
     'cómo se gestiona un incidente y la 86 explica cómo preservar la evidencia mientras eso ocurre.')
make_table(
    ['Año', 'Hito'],
    [['1901', 'Se funda el National Bureau of Standards.'],
     ['1988', 'Pasa a denominarse NIST y amplía su mandato hacia la innovación tecnológica.'],
     ['Década de 1990', 'Se consolida la serie Special Publications 800 dedicada a seguridad informática.'],
     ['2006', 'Se publica el SP 800-86 sobre integración de técnicas forenses en la respuesta a incidentes.']],
    [3.0, 12.9])
caption('Tabla 1. Hitos institucionales del NIST relevantes para la guía analizada.')

# ══════════════════════════════════════════════════════════════════════════
# 5. QUIÉNES ELABORAN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. ¿Quiénes elaboran la guía?', level=1)
para('El SP 800-86 fue elaborado por la Computer Security Division, perteneciente al Information '
     'Technology Laboratory del NIST. Sus autores son Karen Kent, Suzanne Chevalier, Tim Grance y Hung '
     'Dang, investigadores vinculados a esa división y a la práctica profesional de la respuesta a '
     'incidentes.')
para('El documento se publicó en 2006 y, pese al tiempo transcurrido, continúa vigente y ampliamente '
     'citado. Esto se explica porque su aporte no está en las herramientas concretas —que evidentemente '
     'han cambiado— sino en la estructura metodológica del proceso, que se mantiene válida con '
     'independencia de la tecnología del momento.')
make_table(
    ['Dato', 'Detalle'],
    [['Título', 'Guide to Integrating Forensic Techniques into Incident Response'],
     ['Identificador', 'NIST Special Publication 800-86'],
     ['Organismo', 'National Institute of Standards and Technology (NIST)'],
     ['Unidad responsable', 'Computer Security Division – Information Technology Laboratory'],
     ['Autores', 'Karen Kent, Suzanne Chevalier, Tim Grance y Hung Dang'],
     ['Año de publicación', '2006'],
     ['Disponibilidad', 'Descarga libre y gratuita en csrc.nist.gov'],
     ['Estado', 'Vigente; referencia internacional en forense digital']],
    [4.4, 11.5])
caption('Tabla 2. Ficha técnica de la publicación.')

# ══════════════════════════════════════════════════════════════════════════
# 6. ALCANCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Alcance de la guía', level=1)
para('La guía se dirige a quienes deben manejar evidencia digital en el marco de un incidente de '
     'seguridad, y no exclusivamente a peritos judiciales. Este es un rasgo característico del documento: '
     'reconoce que, en la práctica, la primera persona que toca la evidencia rara vez es un forense '
     'profesional, sino el administrador de sistemas o el analista que detectó el problema.')
bullets([
    ('Equipos de respuesta a incidentes.', 'Personal que atiende el evento y necesita recolectar y '
     'preservar la evidencia sin destruirla durante las tareas de contención y erradicación.'),
    ('Investigadores forenses digitales.', 'Peritos que analizan medios de almacenamiento, memoria, '
     'tráfico de red y aplicaciones para sustentar técnicamente una investigación.'),
    ('Administradores de sistemas y redes.', 'Técnicos que operan la infraestructura y deben conservar '
     'registros y evidencia ante un evento de seguridad.'),
    ('Organizaciones públicas y privadas.', 'Instituciones que necesitan definir políticas y '
     'procedimientos forenses propios antes de que ocurra el incidente.'),
])
para('En cuanto a la materia que cubre, el documento aborda el proceso forense general y luego lo '
     'particulariza según el origen de los datos: archivos, sistemas operativos, tráfico de red y '
     'aplicaciones. Dedica también un capítulo a la construcción de una capacidad forense '
     'organizacional, es decir, a las políticas, los roles, la capacitación y las herramientas que una '
     'institución debería tener resueltas antes de sufrir un incidente.')
para('Igual de importante es reconocer sus límites, que el propio NIST declara: la guía no constituye '
     'asesoría legal, no reemplaza la normativa procesal de cada país y no es el manual de ninguna '
     'herramienta en particular. En el caso boliviano, esto significa que el proceso metodológico del '
     'SP 800-86 debe articularse con lo que establece la legislación nacional en materia de prueba '
     'digital y peritaje.')

# ══════════════════════════════════════════════════════════════════════════
# 7. PANORAMA COMPARATIVO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Panorama de las guías del área forense', level=1)
para('Antes de entrar en el detalle del SP 800-86 conviene situarlo entre las demás referencias del área. '
     'La tabla siguiente resume las seis guías estudiadas en la materia.')
make_table(
    ['Guía', 'Origen', 'Enfoque principal', 'Año'],
    [['RFC 3227', 'IETF', 'Recolección y archivo de evidencia; establece el orden de volatilidad', '2002'],
     ['SWGDE', 'Estados Unidos', 'Buenas prácticas y control de calidad en laboratorios forenses', '1998 en adelante'],
     ['ACPO', 'Reino Unido', 'Cuatro principios de buena práctica para el ámbito policial', '2012 (v5)'],
     ['CP4DF', 'Académico (Alemania)', 'Modelo de proceso común entre respuesta a incidentes y pericia forense', '2007'],
     ['NIST SP 800-86', 'NIST – Estados Unidos', 'Proceso de cuatro fases integrado a la respuesta a incidentes', '2006'],
     ['ISO/IEC 27037', 'Internacional', 'Identificación, recolección, adquisición y preservación de la evidencia', '2012']],
    [3.0, 3.3, 7.2, 2.4], highlight=4)
caption('Tabla 3. Comparación de las principales guías del área forense digital.')
para('La comparación deja ver que las guías no compiten entre sí, sino que se complementan. El RFC 3227 '
     'aporta un criterio operativo insustituible —el orden de volatilidad—, la ISO/IEC 27037 formaliza la '
     'preservación con vocación internacional, el ACPO fija principios de conducta y el SWGDE se ocupa de '
     'la calidad del laboratorio. El SP 800-86 es la que articula todos esos elementos dentro de una '
     'secuencia de trabajo, y por eso resulta la más adecuada como eje de un estudio.')

# ══════════════════════════════════════════════════════════════════════════
# 8. EL PROCESO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. El proceso forense según el NIST SP 800-86', level=1)
doc.add_heading('8.1 Visión general del proceso', level=2)
para('El corazón de la guía es un proceso de cuatro fases: recolección, examen, análisis y reporte. Su '
     'valor no reside únicamente en la división en etapas —otras guías también lo hacen— sino en que '
     'define qué recibe y qué entrega cada una. El SP 800-86 lo expresa como una cadena de '
     'transformación progresiva:')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
r = p.add_run('MEDIOS  →  DATOS  →  INFORMACIÓN  →  EVIDENCIA')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT; r.font.name = BODY
para('Los medios son los soportes físicos y lógicos donde reside la información. La recolección los '
     'convierte en datos adquiridos de forma íntegra. El examen convierte esos datos en información '
     'pertinente al caso, descartando el enorme volumen de material irrelevante. El análisis convierte '
     'esa información en evidencia, es decir, en afirmaciones sostenidas sobre lo que ocurrió. Y el '
     'reporte convierte la evidencia en un documento comunicable y defendible. Saltarse una fase no '
     'ahorra tiempo: simplemente traslada el problema a la siguiente, que trabajará sobre material que no '
     'puede sostener.')
make_table(
    ['Fase', 'Recibe', 'Entrega'],
    [['1. Recolección (Collection)', 'Medios y fuentes de datos', 'Datos adquiridos con integridad verificada'],
     ['2. Examen (Examination)', 'Datos adquiridos', 'Información relevante extraída y clasificada'],
     ['3. Análisis (Analysis)', 'Información relevante', 'Evidencia y conclusiones sustentadas'],
     ['4. Reporte (Reporting)', 'Evidencia y conclusiones', 'Informe técnico, ejecutivo y soporte legal']],
    [5.3, 5.3, 5.3])
caption('Tabla 4. Entradas y salidas de cada fase del proceso.')
para('Las fases son secuenciales pero no rígidas. En la práctica es habitual volver de la fase de '
     'análisis a la de examen cuando un hallazgo obliga a buscar datos que en su momento se habían '
     'descartado; lo que nunca puede hacerse es retroceder a la recolección sobre una fuente ya alterada.')

PHASES = [
    dict(
        h='8.2 Fase 1: Recolección (Collection)',
        intro='La primera fase consiste en identificar, etiquetar, registrar y adquirir los datos de las '
              'fuentes relevantes, preservando su integridad. Es la fase más delicada del proceso, porque '
              'los errores que se cometen aquí son irreversibles: un dato sobrescrito no se recupera y una '
              'evidencia mal adquirida no se puede volver a adquirir bien.',
        sub=[
            ('Identificación de las fuentes',
             'Se determinan los medios físicos involucrados (discos duros, unidades de estado sólido, '
             'memorias USB, teléfonos), la memoria RAM del equipo, los registros del sistema operativo, el '
             'tráfico de red y los datos generados por aplicaciones. La priorización sigue el orden de '
             'volatilidad: lo que se pierde antes se captura antes.'),
            ('Adquisición forense',
             'Se realiza una copia bit a bit del medio, empleando un bloqueador de escritura que impida '
             'cualquier modificación del original. La regla es absoluta: nunca se trabaja sobre la fuente '
             'original, sino sobre la imagen. Cada imagen se sella con funciones hash (MD5 y SHA-256) que '
             'permitirán demostrar después que el contenido no cambió.'),
            ('Cadena de custodia',
             'Se abre el registro formal e ininterrumpido de quién tuvo la evidencia, en qué momento, con '
             'qué finalidad y bajo qué condiciones de resguardo. Un vacío en este registro es suficiente '
             'para que la defensa cuestione la totalidad del material recolectado.'),
        ],
        tools='dd y dcfldd, FTK Imager, bloqueadores de escritura por hardware, Volatility para captura y '
              'análisis de memoria, Wireshark y tcpdump para tráfico de red.',
        out='Imagen forense verificada mediante hash y formulario de cadena de custodia firmado.',
    ),
    dict(
        h='8.3 Fase 2: Examen (Examination)',
        intro='La segunda fase procesa los datos adquiridos —de forma automatizada y manual— para hacer '
              'visible y manejable lo que interesa al caso. El problema que resuelve es de escala: un disco '
              'de un terabyte contiene millones de archivos, de los cuales quizá una docena sean relevantes.',
        sub=[
            ('Reducción del volumen',
             'Se descartan los archivos conocidos del sistema operativo y de las aplicaciones comparando '
             'sus hashes contra listas de referencia como la NSRL del propio NIST. Este filtrado por listas '
             'blancas reduce drásticamente el universo de análisis sin perder información relevante.'),
            ('Recuperación de datos',
             'Se recuperan archivos borrados, particiones ocultas, contenido del espacio no asignado y del '
             'slack space. También se descomprimen contenedores, se descifra el material cuando se dispone '
             'legítimamente de la clave y se recuperan archivos por firma mediante técnicas de carving.'),
            ('Filtrado y clasificación',
             'Se busca por palabras clave, tipo de archivo, rango de fechas, usuario propietario o '
             'extensiones falsificadas. El resultado se organiza en un conjunto acotado y ordenado que la '
             'fase de análisis pueda interpretar.'),
        ],
        tools='Autopsy y The Sleuth Kit, EnCase, bulk_extractor, PhotoRec y Foremost para carving.',
        out='Conjunto acotado de datos relevantes, extraídos y clasificados, listo para el análisis.',
    ),
    dict(
        h='8.4 Fase 3: Análisis (Analysis)',
        intro='La tercera fase correlaciona la información obtenida y deriva conclusiones que respondan a '
              'las preguntas del caso: qué ocurrió, cuándo empezó, por dónde entró el atacante, hasta dónde '
              'llegó y qué se llevó. Es la fase que exige más criterio profesional, porque las herramientas '
              'entregan datos, pero la interpretación es humana.',
        sub=[
            ('Construcción de la línea de tiempo',
             'Se reconstruye la cronología de accesos, creaciones, modificaciones y borrados a partir de '
             'las marcas temporales del sistema de archivos, los registros de eventos y los artefactos de '
             'aplicación. La línea de tiempo es el entregable central del análisis.'),
            ('Correlación entre fuentes',
             'Se cruzan los registros de red con los del sistema operativo y los de las aplicaciones. Un '
             'indicio aislado rara vez prueba algo; la fuerza del análisis está en que varias fuentes '
             'independientes coincidan en la misma versión de los hechos.'),
            ('Identificación de artefactos',
             'Se detectan el malware y sus mecanismos de persistencia, las cuentas comprometidas, las '
             'conexiones no autorizadas, los datos exfiltrados y los rastros de técnicas antiforenses, como '
             'el borrado selectivo de registros o la manipulación de marcas temporales.'),
        ],
        tools='log2timeline y Plaso para líneas de tiempo, Volatility para memoria, NetworkMiner para '
              'tráfico, YARA para identificación de patrones y RegRipper para el registro de Windows.',
        out='Hipótesis confirmada o descartada, respaldada por una línea de tiempo y los artefactos que la '
            'sustentan.',
    ),
    dict(
        h='8.5 Fase 4: Reporte (Reporting)',
        intro='La última fase documenta el proceso seguido, los hallazgos obtenidos y las conclusiones '
              'alcanzadas. Suele ser la fase menos valorada por los técnicos y, sin embargo, es la que '
              'determina si el trabajo sirve para algo: una investigación impecable mal documentada es una '
              'investigación perdida.',
        sub=[
            ('Informe técnico',
             'Detalla la metodología aplicada, las herramientas y versiones utilizadas, los comandos '
             'ejecutados, los hashes calculados y los hallazgos. El criterio de calidad es la '
             'reproducibilidad: otro perito debe poder repetir el procedimiento y llegar al mismo '
             'resultado.'),
            ('Informe ejecutivo',
             'Resume el incidente sin lenguaje técnico para la gerencia y las áreas no técnicas: qué pasó, '
             'cuándo empezó, qué resultó afectado, cuál fue la causa raíz y qué se recomienda hacer. Es la '
             'parte del trabajo que efectivamente se lee al momento de tomar decisiones.'),
            ('Soporte legal',
             'Incluye la cadena de custodia completa, la bitácora de cada acción realizada y los anexos con '
             'los hashes de la evidencia. Es el material que sostiene la admisibilidad de lo actuado ante '
             'la autoridad competente.'),
        ],
        tools='Módulos de reporte de Autopsy, CaseNotes para bitácora, plantillas institucionales de cadena '
              'de custodia y firma digital del informe final.',
        out='Informe técnico, informe ejecutivo, cadena de custodia documentada y recomendaciones de '
            'remediación.',
    ),
]
for ph in PHASES:
    doc.add_heading(ph['h'], level=2)
    para(ph['intro'])
    for title, body in ph['sub']:
        doc.add_heading(title, level=3)
        para(body)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('Herramientas típicas: '); r.bold = True; r.font.name = BODY; r.font.size = Pt(10)
    r = p.add_run(ph['tools']); r.font.name = BODY; r.font.size = Pt(10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('Resultado de la fase: '); r.bold = True; r.font.name = BODY; r.font.size = Pt(10)
    r = p.add_run(ph['out']); r.font.name = BODY; r.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════
# 9. CONSIDERACIONES TRANSVERSALES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Consideraciones transversales al proceso', level=1)
para('Hay tres elementos que no pertenecen a una fase en particular, sino que atraviesan las cuatro y '
     'condicionan la validez de todo el trabajo.')
doc.add_heading('9.1 Cadena de custodia', level=2)
para('Es el registro documental que acredita el recorrido íntegro de la evidencia desde su recolección '
     'hasta su presentación. Debe consignar quién la recibió, cuándo, de quién, con qué finalidad y bajo '
     'qué condiciones de almacenamiento. Su función es demostrar que el material presentado es el mismo '
     'que se recolectó y que nadie pudo alterarlo en el trayecto. Un solo tramo sin documentar basta para '
     'poner en duda el conjunto.')
doc.add_heading('9.2 Orden de volatilidad', level=2)
para('Los datos desaparecen a velocidades muy distintas. El criterio, tomado del RFC 3227 y recogido por '
     'el NIST, indica capturar primero lo que se pierde antes. El orden habitual, de más volátil a menos, '
     'es el siguiente:')
bullets([
    'Registros del procesador y memoria caché.',
    'Tabla de rutas, caché ARP, tabla de procesos y estadísticas del núcleo.',
    'Memoria RAM y conexiones de red activas.',
    'Archivos temporales del sistema.',
    'Contenido del disco duro y de los medios de almacenamiento.',
    'Registros y monitoreo almacenados de forma remota.',
    'Configuración física y topología de la red.',
    'Medios de respaldo y archivos históricos.',
], style='List Number')
doc.add_heading('9.3 Admisibilidad y aspectos legales', level=2)
para('El SP 800-86 insiste en que el criterio técnico no basta. La evidencia debe haberse obtenido de '
     'forma lícita, con la autorización correspondiente, y su tratamiento debe respetar la normativa de '
     'protección de datos y de privacidad aplicable. Por eso la guía recomienda involucrar al área legal '
     'desde el primer momento del incidente y no cuando el informe ya está escrito. En el contexto '
     'boliviano, esta articulación con el marco procesal penal es indispensable si se pretende que el '
     'material sirva como prueba.')

# ══════════════════════════════════════════════════════════════════════════
# 10. EJEMPLO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Ejemplo de aplicación: incidente de ransomware', level=1)
doc.add_heading('10.1 Planteamiento del caso', level=2)
para('Se plantea el siguiente caso hipotético. Una empresa mediana inicia la jornada y descubre que los '
     'archivos de varios servidores se encuentran cifrados y que se ha depositado una nota exigiendo un '
     'rescate. La sospecha inicial apunta a un ransomware que habría ingresado por correo electrónico en '
     'el equipo del área de administración. La gerencia solicita determinar cómo entró el atacante, hasta '
     'dónde llegó y si hubo salida de información antes del cifrado.')
doc.add_heading('10.2 Fase 1: Recolección', level=2)
bullets([
    'Se aísla físicamente el equipo sospechoso de la red, sin apagarlo, para no perder la memoria volátil.',
    'Se realiza el volcado de la memoria RAM antes que cualquier otra acción, respetando el orden de volatilidad.',
    'Se apaga el equipo y se extrae el disco, del que se obtiene una imagen bit a bit con bloqueador de escritura.',
    'Se calcula el hash SHA-256 de la imagen y se registra en el acta de cadena de custodia, firmada por el '
    'responsable de TI y el investigador.',
    'Se solicitan y preservan los registros del servidor de correo, del cortafuegos y del servidor proxy '
    'correspondientes a las últimas setenta y dos horas.',
])
doc.add_heading('10.3 Fase 2: Examen', level=2)
bullets([
    'Se monta la imagen en modo solo lectura sobre la estación forense.',
    'Se descartan los archivos conocidos del sistema operativo mediante comparación de hashes contra la NSRL.',
    'Se recupera desde el espacio no asignado el archivo adjunto que el usuario había eliminado.',
    'Se listan los ejecutables creados o modificados en la ventana temporal del incidente.',
    'Se identifican los archivos cifrados y la nota de rescate, y se extraen sus marcas temporales.',
])
doc.add_heading('10.4 Fase 3: Análisis', level=2)
para('Del cruce entre el volcado de memoria, la imagen del disco y los registros de red se reconstruye la '
     'siguiente secuencia de hechos. Cada afirmación se acompaña del artefacto que la respalda, que es lo '
     'que la convierte en evidencia y no en suposición.')
make_table(
    ['Hora', 'Evento reconstruido', 'Artefacto que lo prueba'],
    [['09:14', 'Ingresa un correo con el adjunto Factura_0725.xlsm', 'Registros del servidor de correo'],
     ['09:22', 'El usuario abre el archivo y habilita las macros', 'Claves TrustRecords del registro y archivo Prefetch'],
     ['09:23', 'Se descarga el ejecutable desde una dirección IP externa', 'Captura de red y registros del servidor proxy'],
     ['09:25', 'Se establece persistencia mediante clave Run y tarea programada', 'Registro de Windows y carpeta Tasks'],
     ['10:05', 'Movimiento lateral hacia dos servidores mediante SMB', 'Eventos 4624 y 4672 del registro de seguridad'],
     ['11:40', 'Exfiltración de datos hacia un servicio de almacenamiento externo', 'Registros NetFlow y del cortafuegos'],
     ['12:10', 'Cifrado masivo de archivos y creación de la nota de rescate', 'Marcas temporales de la MFT del sistema de archivos']],
    [1.6, 7.2, 7.1])
caption('Tabla 5. Línea de tiempo reconstruida del incidente y su respaldo probatorio.')
para('El análisis permite concluir que el vector de entrada fue una macro contenida en un archivo adjunto, '
     'que el atacante permaneció aproximadamente tres horas dentro de la red antes de ejecutar el cifrado '
     'y —dato crítico para la organización— que hubo exfiltración de información antes del cifrado, lo '
     'que convierte el hecho en una fuga de datos y no solamente en una interrupción del servicio.')
doc.add_heading('10.5 Fase 4: Reporte', level=2)
bullets([
    ('Informe técnico.', 'Metodología, herramientas y versiones, comandos ejecutados, hashes de cada '
     'evidencia y hallazgos detallados con sus artefactos de respaldo.'),
    ('Informe ejecutivo.', 'Resumen del incidente, línea de tiempo simplificada, alcance del daño, '
     'confirmación de la fuga de información e impacto estimado para la operación.'),
    ('Anexo legal.', 'Cadena de custodia completa, bitácora de acciones y listado de hashes, preparado '
     'para su eventual presentación ante la autoridad competente.'),
    ('Plan de remediación.', 'Bloqueo de macros por directiva, segmentación de la red, revisión de '
     'credenciales de administrador local, restauración desde respaldos verificados y capacitación del '
     'personal en reconocimiento de correos maliciosos.'),
])

# ══════════════════════════════════════════════════════════════════════════
# 11. RECOMENDACIONES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Recomendaciones', level=1)
bullets([
    ('Preparar la capacidad forense antes del incidente.', 'El propio NIST dedica un capítulo a este '
     'punto: las políticas, los roles, las herramientas y la capacitación deben estar resueltos con '
     'anticipación. Improvisar durante el incidente es la causa más frecuente de pérdida de evidencia.'),
    ('No trabajar nunca sobre la evidencia original.', 'Toda operación se realiza sobre copias forenses '
     'verificadas, obtenidas con bloqueador de escritura.'),
    ('Verificar la integridad en cada paso.', 'Calcular y comparar los hashes al inicio y al final de cada '
     'fase. Un valor distinto rompe la cadena de evidencia y obliga a documentar la causa.'),
    ('Respetar el orden de volatilidad.', 'Capturar la memoria y el estado de la red antes de apagar el '
     'equipo. Apagar primero y pensar después destruye evidencia irrecuperable.'),
    ('Documentar cada acción en el momento en que se realiza.', 'Quién, cuándo, con qué herramienta y con '
     'qué parámetros. Lo que no está documentado, a efectos prácticos, no ocurrió.'),
    ('Emplear herramientas reconocidas y registrar su versión.', 'La validez del resultado depende en '
     'parte de la trazabilidad del instrumento utilizado.'),
    ('Involucrar al área legal desde el inicio.', 'La licitud de la obtención condiciona la admisibilidad '
     'del material, y esa decisión no es técnica.'),
    ('Mantener la capacitación del equipo.', 'Los sistemas de archivos, los mecanismos de cifrado y las '
     'técnicas antiforenses cambian permanentemente; una metodología sólida no compensa el '
     'desconocimiento técnico actualizado.'),
])

# ══════════════════════════════════════════════════════════════════════════
# 12. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Conclusiones', level=1)
bullets([
    'El NIST SP 800-86 convierte la informática forense en un proceso repetible y auditable de cuatro '
    'fases, en lugar de un conjunto de técnicas aplicadas según el criterio del momento. Esa estructura '
    'es precisamente lo que permite sostener un resultado ante terceros.',
    'Su aporte distintivo frente a otras guías es integrar la práctica forense dentro de la respuesta a '
    'incidentes: la evidencia se preserva mientras se contiene el ataque, y no como una actividad '
    'posterior cuando buena parte del rastro ya se perdió.',
    'La cadena medios – datos – información – evidencia hace explícito el valor que agrega cada etapa y '
    'evita la tentación de saltar pasos por urgencia, error que traslada el problema a la fase siguiente '
    'en lugar de resolverlo.',
    'La guía es complementaria y no excluyente respecto de las demás referencias del área: se apoya en el '
    'orden de volatilidad del RFC 3227, converge con los criterios de preservación de la ISO/IEC 27037 y '
    'no contradice los principios del ACPO.',
    'El ejemplo desarrollado demuestra que el modelo es aplicable con herramientas de software libre y '
    'sin infraestructura costosa, lo que lo hace viable para organizaciones medianas y para el ámbito '
    'académico.',
    'Por su claridad metodológica, su gratuidad y su aplicabilidad inmediata, el SP 800-86 resulta la '
    'guía más didáctica de las seis analizadas para introducirse a la informática forense, sin que ello '
    'implique afirmar su superioridad técnica sobre las demás.',
], style='List Number')

# ══════════════════════════════════════════════════════════════════════════
# 13. BIBLIOGRAFÍA
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Bibliografía', level=1)
refs = [
    'Kent, K., Chevalier, S., Grance, T. y Dang, H. (2006). Guide to Integrating Forensic Techniques into '
    'Incident Response (NIST Special Publication 800-86). Gaithersburg: National Institute of Standards '
    'and Technology.',
    'Cichonski, P., Millar, T., Grance, T. y Scarfone, K. (2012). Computer Security Incident Handling '
    'Guide (NIST Special Publication 800-61 Rev. 2). Gaithersburg: National Institute of Standards and '
    'Technology.',
    'Brezinski, D. y Killalea, T. (2002). Guidelines for Evidence Collection and Archiving (RFC 3227, '
    'BCP 55). Internet Engineering Task Force.',
    'International Organization for Standardization (2012). ISO/IEC 27037: Information technology – '
    'Security techniques – Guidelines for identification, collection, acquisition and preservation of '
    'digital evidence. Ginebra: ISO.',
    'Association of Chief Police Officers (2012). ACPO Good Practice Guide for Digital Evidence, '
    'versión 5. Reino Unido: ACPO.',
    'Freiling, F. C. y Schwittay, B. (2007). A Common Process Model for Incident Response and Computer '
    'Forensics. En Proceedings of the IT Incident Management and IT Forensics (IMF). Alemania.',
    'Scientific Working Group on Digital Evidence (2018). SWGDE Best Practices for Computer Forensic '
    'Acquisitions. Estados Unidos: SWGDE.',
]
for r_ in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(r_); run.font.name = BODY; run.font.size = Pt(10.5)

doc.sections[0].different_first_page_header_footer = True
add_page_number_footer(doc.sections[0])
doc.save('NIST_SP_800-86_Documento.docx')
print('OK — NIST_SP_800-86_Documento.docx')
