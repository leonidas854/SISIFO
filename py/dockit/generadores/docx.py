"""Genera un .docx académico a partir de un guion.

Estilos sobrios de trabajo universitario: cuerpo justificado, títulos
jerarquizados, tablas con cabecera, figuras con leyenda y su fuente. Se abre
igual en OnlyOffice, Word y LibreOffice.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import campos
from . import guion as G

TINTA = RGBColor(0x1A, 0x1A, 0x1A)
SUAVE = RGBColor(0x5A, 0x5A, 0x5A)
PALABRAS_POR_PAGINA = 450  # para estimar; la medida real la da LibreOffice


def _estilos(doc: Document, tipografia: str, cuerpo_pt: float,
             interlineado: float = 1.5) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = tipografia
    normal.font.size = Pt(cuerpo_pt)
    normal.font.color.rgb = TINTA
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), tipografia)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = interlineado  # APA 7 pide doble; el BRIEF lo decide
    pf.space_after = Pt(6)

    for nombre, tam, antes, despues in [
        ("Heading 1", cuerpo_pt + 5, 18, 8),
        ("Heading 2", cuerpo_pt + 2, 14, 6),
        ("Heading 3", cuerpo_pt + 1, 10, 4),
        ("Heading 4", cuerpo_pt, 8, 4),
    ]:
        if nombre not in doc.styles:
            continue
        e = doc.styles[nombre]
        e.font.name = tipografia
        e.font.size = Pt(tam)
        e.font.bold = True
        e.font.color.rgb = TINTA
        e.paragraph_format.space_before = Pt(antes)
        e.paragraph_format.space_after = Pt(despues)
        e.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _numeracion_de_pagina(doc: Document) -> None:
    """Número de página centrado en el pie, como campo de Word.

    Se construye con el mismo helper que los índices: antes se creaba un
    elemento inventado `<w:fldbegin>` y el pie acababa mostrando la palabra
    «PAGE» en lugar del número.
    """
    pie = doc.sections[0].footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    campos._run_con_campo(pie, "PAGE", "1")


def _portada(doc: Document, guion: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(guion["titulo"])
    r.bold = True
    r.font.size = Pt(20)
    if guion.get("autor"):
        a = doc.add_paragraph()
        a.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = a.add_run(guion["autor"])
        ra.font.size = Pt(12)
        ra.font.color.rgb = SUAVE
    doc.add_page_break()


def _tabla(doc: Document, b: dict, tipografia: str) -> None:
    cab = b.get("cabecera") or []
    filas = b.get("filas") or []
    ncol = len(cab) if cab else len(filas[0])
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    if cab:
        fila = t.add_row().cells
        for i, texto in enumerate(cab):
            fila[i].text = ""
            run = fila[i].paragraphs[0].add_run(str(texto))
            run.bold = True
            run.font.name = tipografia
    for f in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(f):
            celdas[i].text = ""
            celdas[i].paragraphs[0].add_run(str(valor)).font.name = tipografia

    if b.get("leyenda"):
        campos.leyenda(doc, campos.ROTULO_TABLA, b["leyenda"], b.get("fuente", ""),
                       WD_ALIGN_PARAGRAPH.LEFT)


def _figura(doc: Document, b: dict) -> None:
    ruta = Path(b["ruta"])
    if ruta.exists():
        doc.add_picture(str(ruta), width=Cm(14))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if b.get("leyenda") or b.get("fuente"):
        campos.leyenda(doc, campos.ROTULO_FIGURA, b.get("leyenda", ""),
                       b.get("fuente", ""), WD_ALIGN_PARAGRAPH.CENTER)


def generar(guion: dict, destino: str, bibliografia: dict[str, str],
            en_texto: dict[str, str], formato: dict | None = None) -> dict:
    G.validar(guion, set(bibliografia) if bibliografia else None)
    formato = formato or {}
    tipografia = formato.get("tipografia", "Calibri")
    cuerpo_pt = float(formato.get("tamano_pt", 11))
    margen = float(formato.get("margenes_cm", 2.5))

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(margen)
    sec.left_margin = sec.right_margin = Cm(margen)
    _estilos(doc, tipografia, cuerpo_pt,
             float(formato.get("interlineado", 1.5)))
    _numeracion_de_pagina(doc)
    _portada(doc, guion)

    palabras = 0
    for b in guion["bloques"]:
        clase = b["clase"]
        if clase == "titulo":
            doc.add_heading(b["texto"], level=min(b.get("nivel", 1), 4))
            palabras += len(b["texto"].split())
        elif clase == "parrafo":
            texto = G.texto_con_citas(b, en_texto)
            doc.add_paragraph(texto)
            palabras += len(texto.split())
        elif clase == "cita":
            p = doc.add_paragraph(G.texto_con_citas(b, en_texto))
            p.paragraph_format.left_indent = Cm(1.27)   # sangría APA
            p.paragraph_format.line_spacing = 1.0
            palabras += len(b["texto"].split())
        elif clase == "lista":
            for it in b["items"]:
                doc.add_paragraph(str(it), style="List Bullet")
                palabras += len(str(it).split())
        elif clase == "tabla":
            _tabla(doc, b, tipografia)
            palabras += sum(len(str(c).split()) for f in b["filas"] for c in f)
        elif clase == "figura":
            _figura(doc, b)
        elif clase == "indice":
            campos.indice_contenido(doc, b.get("texto") or "Índice",
                                    b.get("notas") or "1-3")
            doc.add_page_break()
        elif clase == "indice_tablas":
            campos.indice_de(doc, campos.ROTULO_TABLA,
                             b.get("texto") or "Índice de tablas")
            doc.add_page_break()
        elif clase == "indice_figuras":
            campos.indice_de(doc, campos.ROTULO_FIGURA,
                             b.get("texto") or "Índice de ilustraciones")
            doc.add_page_break()
        elif clase == "salto":
            doc.add_page_break()
        elif clase == "bibliografia":
            doc.add_page_break()
            doc.add_heading("Referencias", level=1)
            for entrada in sorted(bibliografia.values()):
                p = doc.add_paragraph(entrada)
                # sangría francesa, como pide APA
                p.paragraph_format.first_line_indent = Cm(-1.27)
                p.paragraph_format.left_indent = Cm(1.27)
                palabras += len(entrada.split())

    campos.actualizar_campos_al_abrir(doc)
    Path(destino).parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)
    return {"ruta": destino,
            "unidades": max(1, round(palabras / PALABRAS_POR_PAGINA) + 1),
            "unidades_estimadas": True}
