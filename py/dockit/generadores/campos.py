"""Campos de Word: índices automáticos y leyendas numeradas.

Un trabajo de tesis exige que el índice de contenido, el de tablas y el de
ilustraciones **no se escriban a mano**. En OOXML eso son campos:

- `TOC \\o "1-3"` recorre los estilos Heading y arma el índice de contenido.
- `TOC \\c "Tabla"` recoge las leyendas numeradas con `SEQ Tabla`.
- `SEQ Tabla \\* ARABIC` numera cada leyenda sola; si se escribe el número a
  mano, el índice no la ve y la numeración se descuadra al insertar una tabla.

Los campos se calculan al abrir el documento, así que además se marca
`w:updateFields` en settings.xml. Sin eso, el usuario abre el archivo y ve
«Actualice el índice» en lugar del índice.
"""
from __future__ import annotations

from docx.document import Document as _Doc
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROTULO_TABLA = "Tabla"
ROTULO_FIGURA = "Ilustración"   # el rótulo que espera Word en español


def _run_con_campo(parrafo: Paragraph, instruccion: str,
                   marcador: str = "") -> None:
    """Inserta un campo completo: begin → instrText → separate → texto → end.

    El `marcador` es lo que se ve mientras el campo no está calculado. Con un
    texto claro, quien abra el documento sabe que debe pulsar F9 si su lector
    no actualiza solo.
    """
    r = parrafo.add_run()
    ini = OxmlElement("w:fldChar")
    ini.set(qn("w:fldCharType"), "begin")
    ini.set(qn("w:dirty"), "true")          # obliga a recalcular al abrir
    r._r.append(ini)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruccion} "
    r._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)

    if marcador:
        t = OxmlElement("w:t")
        t.text = marcador
        r._r.append(t)

    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    r._r.append(fin)


def indice_contenido(doc: _Doc, titulo: str = "Índice", niveles: str = "1-3") -> None:
    """Índice de contenido a partir de los estilos Heading."""
    if titulo:
        doc.add_heading(titulo, level=1)
    p = doc.add_paragraph()
    _run_con_campo(p, f'TOC \\o "{niveles}" \\h \\z \\u',
                   "Índice de contenido — pulse F9 para actualizar")


def indice_de(doc: _Doc, rotulo: str, titulo: str) -> None:
    """Índice de tablas o de ilustraciones, según el rótulo de las leyendas."""
    if titulo:
        doc.add_heading(titulo, level=1)
    p = doc.add_paragraph()
    _run_con_campo(p, f'TOC \\h \\z \\c "{rotulo}"',
                   f"Índice de {titulo.lower()} — pulse F9 para actualizar")


def leyenda(doc: _Doc, rotulo: str, texto: str, fuente: str = "",
            alineacion=None) -> Paragraph:
    """Leyenda numerada con SEQ, con el estilo Caption que el índice recoge.

    Queda «Tabla 1. Comparación de enfoques», donde el 1 lo pone Word.
    """
    p = doc.add_paragraph(style="Caption")
    if alineacion is not None:
        p.alignment = alineacion

    corrida = p.add_run(f"{rotulo} ")
    corrida.bold = True
    _run_con_campo(p, f"SEQ {rotulo} \\* ARABIC", "1")

    resto = texto.strip()
    if resto:
        p.add_run(f". {resto}")
    if fuente:
        nota = p.add_run(f" Fuente: {fuente}")
        nota.italic = True
        nota.font.size = Pt(9)
    return p


def actualizar_campos_al_abrir(doc: _Doc) -> None:
    """Marca el documento para que Word/OnlyOffice calculen los índices solos."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)
